import os
import sys
import subprocess
import whisper
from docx import Document
from docx.shared import Pt
from PyQt5 import QtWidgets, QtGui, QtCore

def get_ffmpeg_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(os.path.dirname(sys.executable), 'ffmpeg.exe')
    else:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ffmpeg.exe')

def extract_audio(input_video_file, output_audio_file):
    ffmpeg_path = get_ffmpeg_path()
    subprocess.run([ffmpeg_path, '-y', '-i', input_video_file, '-vn', '-acodec', 'copy', output_audio_file], check=True)

def save_to_docx(transcription, output_file, heading):
    doc = Document()
    doc.add_heading(heading, level=1)
    paragraph = doc.add_paragraph(transcription)
    run = paragraph.runs[0]
    font = run.font
    font.size = Pt(12)
    doc.save(output_file)

class TranscriptionApp(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Transcrição de Vídeo para DOCX')
        self.setGeometry(100, 100, 600, 400)
        self.setStyleSheet("background-color: #2E2E2E;")

        layout = QtWidgets.QVBoxLayout()

        title_label = QtWidgets.QLabel('Video - Docx')
        title_label.setFont(QtGui.QFont('Arial', 18, QtGui.QFont.Bold))
        title_label.setAlignment(QtCore.Qt.AlignCenter)
        title_label.setStyleSheet("color: #FFFFFF;")
        layout.addWidget(title_label)

        video_label = QtWidgets.QLabel('Arquivo de Vídeo:')
        video_label.setStyleSheet("color: #FFFFFF;")
        layout.addWidget(video_label)

        self.video_file_entry = QtWidgets.QLineEdit()
        layout.addWidget(self.video_file_entry)

        browse_button = QtWidgets.QPushButton('Procurar')
        browse_button.setStyleSheet("background-color: #007ACC; color: white;")
        browse_button.clicked.connect(self.select_video_file)
        layout.addWidget(browse_button)

        save_label = QtWidgets.QLabel('Salvar como (local):')
        save_label.setStyleSheet("color: #FFFFFF;")
        layout.addWidget(save_label)

        self.output_file_entry = QtWidgets.QLineEdit()
        layout.addWidget(self.output_file_entry)

        save_button = QtWidgets.QPushButton('Escolher Local para Salvar')
        save_button.setStyleSheet("background-color: #007ACC; color: white;")
        save_button.clicked.connect(self.select_output_file)
        layout.addWidget(save_button)

        heading_label = QtWidgets.QLabel('Título do Documento:')
        heading_label.setStyleSheet("color: #FFFFFF;")
        layout.addWidget(heading_label)

        self.heading_entry = QtWidgets.QLineEdit()
        layout.addWidget(self.heading_entry)

        process_button = QtWidgets.QPushButton('Processar')
        process_button.setStyleSheet("background-color: #007ACC; color: white;")
        process_button.clicked.connect(self.process_video)
        layout.addWidget(process_button)

        self.setLayout(layout)

    def select_video_file(self):
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Selecionar Arquivo de Vídeo", "", "Video Files (*.mp4 *.avi *.mkv)")
        if file_path:
            self.video_file_entry.setText(file_path)

    def select_output_file(self):
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Salvar Como", "", "Word Documents (*.docx)")
        if file_path:
            self.output_file_entry.setText(file_path)

    def process_video(self):
        input_video_file = self.video_file_entry.text()
        output_docx_file = self.output_file_entry.text()
        output_audio_file = os.path.splitext(output_docx_file)[0] + ".aac"
        heading = self.heading_entry.text() or 'Transcrição'

        try:
            extract_audio(input_video_file, output_audio_file)
            modelo = whisper.load_model("small")
            resposta = modelo.transcribe(output_audio_file, fp16=False)
            save_to_docx(resposta["text"], output_docx_file, heading)
            QtWidgets.QMessageBox.information(self, "Sucesso", f"Transcrição salva em: {output_docx_file}")
        except FileNotFoundError:
            QtWidgets.QMessageBox.critical(self, "Erro", "Arquivo de áudio não encontrado.")
        except whisper.AudioFormatError:
            QtWidgets.QMessageBox.critical(self, "Erro", "Formato de áudio não suportado.")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Erro", f"Erro inesperado: {e}")

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = TranscriptionApp()
    window.show()
    sys.exit(app.exec_())